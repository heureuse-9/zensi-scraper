import csv
import html
import io
import json
import os
import re
import shutil
import subprocess
import time
import urllib.parse
import xml.etree.ElementTree as ET
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from statistics import median
from typing import Any

import requests
import xlsxwriter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36"
    ),
    "Accept-Language": "en-US,en;q=0.9",
}

IG_HEADERS = {
    **HEADERS,
    "X-IG-App-ID": "936619743392459",
    "X-Requested-With": "XMLHttpRequest",
}

DISPLAY_NAMES = {
    "Bryson haywood": "Drakelocksin (Bryson Haywood)",
    "Cheung hoi tung": "Cheung Hoi Tung",
    "Keven zeni": "Keven Zeni",
    "morgan schriner": "Morgan Schriner",
}


@dataclass
class RunConfig:
    campaign_name: str
    start_date: date
    end_date: date
    output_dir: Path
    roster_json: Path | None = None
    creators_csv: Path | None = None
    supplemental_analytics_csvs: list[Path] | None = None
    supplemental_docx: Path | None = None
    google_sheet_csv_url: str | None = None
    scrape_public: bool = True
    verify_instagram: bool = True
    chrome_path: str | None = None
    output_stem: str | None = None
    export_csv: bool = True
    export_xlsx: bool = True
    export_docx: bool = True


def parse_date(value: str | date) -> date:
    if isinstance(value, date):
        return value
    return datetime.strptime(str(value), "%Y-%m-%d").date()


def config_from_json(path: str | Path, require_dates: bool = True) -> RunConfig:
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    start = data.get("start_date")
    end = data.get("end_date")
    if require_dates and (not start or not end):
        raise ValueError("Config must include start_date and end_date for a normal run.")
    return RunConfig(
        campaign_name=data.get("campaign_name", "Zensi"),
        start_date=parse_date(start) if start else date.today(),
        end_date=parse_date(end) if end else date.today(),
        output_dir=Path(data["output_dir"]),
        roster_json=Path(data["roster_json"]) if data.get("roster_json") else None,
        creators_csv=Path(data["creators_csv"]) if data.get("creators_csv") else None,
        supplemental_analytics_csvs=[Path(p) for p in data.get("supplemental_analytics_csvs", [])],
        supplemental_docx=Path(data["supplemental_docx"]) if data.get("supplemental_docx") else None,
        google_sheet_csv_url=data.get("google_sheet_csv_url") or None,
        scrape_public=bool(data.get("scrape_public", True)),
        verify_instagram=bool(data.get("verify_instagram", True)),
        chrome_path=data.get("chrome_path") or None,
        output_stem=data.get("output_stem") or None,
        export_csv=bool(data.get("export_csv", True)),
        export_xlsx=bool(data.get("export_xlsx", True)),
        export_docx=bool(data.get("export_docx", True)),
    )


def weekly_window(run_date: date | None = None, lookback_days: int = 7) -> tuple[date, date]:
    """Return the last complete reporting window.

    If this runs Friday at 8am, it reports the previous Friday through Thursday.
    """
    today = run_date or date.today()
    end = today - timedelta(days=1)
    start = end - timedelta(days=max(lookback_days - 1, 0))
    return start, end


def clean_text(value: Any, limit: int | None = None) -> str:
    text = html.unescape(str(value or "")).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    if limit and len(text) > limit:
        return text[: limit - 1].rstrip() + "..."
    return text


def parse_num(value: Any) -> int | None:
    if value is None:
        return None
    text = str(value).strip().replace(",", "")
    if not text or text.upper() == "N/A":
        return None
    mult = 1
    if text[-1:].lower() == "k":
        mult = 1_000
        text = text[:-1]
    elif text[-1:].lower() == "m":
        mult = 1_000_000
        text = text[:-1]
    try:
        return int(float(text) * mult)
    except ValueError:
        return None


def fmt_num(value: Any) -> str:
    if value is None or value == "":
        return "N/A"
    try:
        return f"{int(round(float(value))):,}"
    except (TypeError, ValueError):
        return str(value)


def fmt_pct(value: float | None) -> str:
    return "N/A" if value is None else f"{value:.2f}%"


def fmt_date(value: date) -> str:
    return value.strftime("%b %d, %Y").replace(" 0", " ")


def safe_div(a: float | int | None, b: float | int | None) -> float | None:
    if a is None or b is None or b == 0:
        return None
    return a / b


def normalize_platform(value: str) -> str:
    text = (value or "").strip().lower()
    if text in {"ig", "instagram"}:
        return "Instagram"
    if text in {"tt", "tiktok", "tik tok"}:
        return "TikTok"
    if text in {"yt", "youtube", "youtube shorts"}:
        return "YouTube"
    return value.strip().title()


def split_handles(value: str | None) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        pieces = value
    else:
        text = str(value)
        pieces = re.split(r"[,;\n]+", text) if re.search(r"https?://", text, re.I) else re.split(r"[,;/\n]+", text)
    handles = []
    for raw in pieces:
        handle = str(raw or "").strip().lstrip("@")
        handle = re.sub(r"^https?://(www\.)?", "", handle, flags=re.I)
        handle = re.sub(r"^(instagram\.com/|tiktok\.com/@|youtube\.com/@)", "", handle, flags=re.I)
        handle = handle.strip("/ ")
        if handle and handle.lower() not in [h.lower() for h in handles]:
            handles.append(handle)
    return handles


def merge_handles(target: dict[str, list[str]], key: str, handles: list[str]) -> None:
    for handle in handles:
        if handle and handle.lower() not in [h.lower() for h in target[key]]:
            target[key].append(handle)


def load_roster_json(roster_json: Path | None) -> dict[str, dict[str, list[str]]]:
    roster: dict[str, dict[str, list[str]]] = {}
    if not roster_json or not roster_json.exists():
        return roster
    data = json.loads(roster_json.read_text(encoding="utf-8"))
    creators = data.get("creators", data if isinstance(data, list) else [])
    for row in creators:
        if not row or row.get("active", True) is False:
            continue
        name = (row.get("name") or row.get("creator") or "").strip()
        if not name:
            continue
        entry = roster.setdefault(name, {"ig": [], "tt": [], "yt": []})
        merge_handles(entry, "ig", split_handles(row.get("instagram") or row.get("ig") or row.get("instagramHandle")))
        merge_handles(entry, "tt", split_handles(row.get("tiktok") or row.get("tt") or row.get("tiktokHandle")))
        merge_handles(entry, "yt", split_handles(row.get("youtube") or row.get("yt") or row.get("youtubeHandle")))
    return roster


def save_roster_json(path: Path, roster: dict[str, dict[str, list[str]]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "creators": [
            {
                "name": creator,
                "active": True,
                "instagram": ["@" + h for h in handles.get("ig", [])],
                "tiktok": ["@" + h for h in handles.get("tt", [])],
                "youtube": ["@" + h for h in handles.get("yt", [])],
            }
            for creator, handles in sorted(roster.items())
        ]
    }
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def load_roster(creators_csv: Path | None, google_sheet_csv_url: str | None = None) -> dict[str, dict[str, list[str]]]:
    roster: dict[str, dict[str, list[str]]] = {}
    if creators_csv and creators_csv.exists():
        with creators_csv.open(newline="", encoding="utf-8-sig") as f:
            for row in csv.DictReader(f):
                name = (row.get("creator") or row.get("Creator Name") or "").strip()
                if not name:
                    continue
                entry = roster.setdefault(name, {"ig": [], "tt": [], "yt": []})
                merge_handles(entry, "ig", split_handles(row.get("instagramHandle") or row.get("IG Handle")))
                merge_handles(entry, "tt", split_handles(row.get("tiktokHandle") or row.get("TikTok Handle")))
                merge_handles(entry, "yt", split_handles(row.get("youtubeHandle") or row.get("YT Handle")))
    if google_sheet_csv_url:
        try:
            response = requests.get(google_sheet_csv_url, timeout=30)
            response.raise_for_status()
            for row in csv.DictReader(io.StringIO(response.text)):
                name = (row.get("Creator Name") or row.get("creator") or "").strip()
                if not name:
                    continue
                entry = roster.setdefault(name, {"ig": [], "tt": [], "yt": []})
                merge_handles(entry, "ig", split_handles(row.get("IG Handle") or row.get("instagramHandle")))
                merge_handles(entry, "tt", split_handles(row.get("TikTok Handle") or row.get("tiktokHandle")))
                merge_handles(entry, "yt", split_handles(row.get("YT Handle") or row.get("youtubeHandle")))
        except Exception:
            pass
    return roster


def load_creator_registry(
    roster_json: Path | None = None,
    creators_csv: Path | None = None,
    google_sheet_csv_url: str | None = None,
) -> dict[str, dict[str, list[str]]]:
    roster = load_roster_json(roster_json)
    imported = load_roster(creators_csv, google_sheet_csv_url)
    for creator, handles in imported.items():
        entry = roster.setdefault(creator, {"ig": [], "tt": [], "yt": []})
        for platform in ["ig", "tt", "yt"]:
            merge_handles(entry, platform, handles.get(platform, []))
    return roster


def category_for(text: str | None) -> str:
    value = (text or "").lower()
    if any(k in value for k in ["study", "student", "school", "college", "grades", "finals", "exam", "homework"]):
        return "Study / academics"
    if any(k in value for k in ["brainrot", "doomscroll", "screen time", "phone", "scroll", "screentime"]):
        return "Screen-time reset"
    if any(k in value for k in ["routine", "habit", "productive", "reset", "lock in", "focus"]):
        return "Productivity routine"
    if any(k in value for k in ["brain", "dopamine", "neuroscience"]):
        return "Cognition / dopamine"
    if any(k in value for k in ["friend", "relationship", "situationship", "drama", "hinge"]):
        return "Relatable social"
    return "General motivation"


def account_url(platform: str, handle: str) -> str:
    handle = (handle or "").strip().lstrip("@")
    if not handle:
        return ""
    if platform == "Instagram":
        return f"https://www.instagram.com/{handle}/"
    if platform == "TikTok":
        return f"https://www.tiktok.com/@{handle}"
    return f"https://www.youtube.com/@{handle}"


def post_url(platform: str, handle: str, post_id: str, fallback_url: str | None = None) -> str:
    if fallback_url:
        return fallback_url
    handle = (handle or "").strip().lstrip("@")
    if platform == "Instagram":
        return f"https://www.instagram.com/reel/{post_id}/"
    if platform == "TikTok":
        return f"https://www.tiktok.com/@{handle}/video/{post_id}"
    return f"https://www.youtube.com/shorts/{post_id}"


def post_id_from_url(url: str) -> str:
    if not url:
        return ""
    patterns = [
        r"/video/(\d+)",
        r"/shorts/([A-Za-z0-9_-]+)",
        r"/(?:reel|p)/([A-Za-z0-9_-]+)",
        r"[?&]v=([A-Za-z0-9_-]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return ""


def handle_from_url(platform: str, url: str) -> str:
    if not url:
        return ""
    if platform == "TikTok":
        match = re.search(r"tiktok\.com/@([^/]+)", url)
    elif platform == "YouTube":
        match = re.search(r"youtube\.com/@([^/]+)", url)
    else:
        match = re.search(r"instagram\.com/([^/]+)/", url)
    return match.group(1) if match else ""


def load_csv_posts(videos_csv: Path | None, start: date, end: date) -> list[dict[str, Any]]:
    if not videos_csv or not videos_csv.exists():
        return []
    posts = []
    with videos_csv.open(newline="", encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            try:
                published = parse_date(row.get("publishDate") or row.get("date") or "")
            except ValueError:
                continue
            if not (start <= published <= end):
                continue
            platform = normalize_platform(row.get("platform") or "")
            url = row.get("videoUrl") or row.get("url") or ""
            post_id = row.get("sourceId") or post_id_from_url(url)
            handle = handle_from_url(platform, url)
            title = row.get("caption") or row.get("title") or post_id or url
            posts.append(
                {
                    "creator": (row.get("creator") or "").strip(),
                    "platform": platform,
                    "handle": handle,
                    "date": published,
                    "post_type": "Reel" if platform == "Instagram" else ("Short" if platform == "YouTube" else "Video"),
                    "post_id": post_id,
                    "title": clean_text(title, 220),
                    "views": parse_num(row.get("views")) or 0,
                    "likes": parse_num(row.get("likes")) or 0,
                    "comments": None if platform == "YouTube" else (parse_num(row.get("comments")) or 0),
                    "saves": parse_num(row.get("saves")) if platform == "TikTok" else None,
                    "shares": parse_num(row.get("shares")) if platform == "TikTok" else None,
                    "category": category_for(title),
                    "url": post_url(platform, handle, post_id, url),
                    "_source": "csv",
                }
            )
    return posts


def canonical_creator(name: str) -> str:
    reverse = {display: canonical for canonical, display in DISPLAY_NAMES.items()}
    return reverse.get(name.strip(), name.strip())


def iter_doc_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def parse_ref(text: str) -> tuple[str, str, str]:
    ref, pid = text, ""
    if " | ID: " in text:
        ref, pid = text.rsplit(" | ID: ", 1)
    caption, handle = ref, ""
    if " | " in ref:
        caption, handle = ref.rsplit(" | ", 1)
    else:
        lines = [line.strip() for line in ref.splitlines() if line.strip()]
        if len(lines) > 1:
            caption, handle = " ".join(lines[:-1]), lines[-1]
    return clean_text(caption, 220), handle.strip().lstrip("@"), pid.strip()


def parse_engagement(text: str, platform: str) -> dict[str, Any]:
    def find(label: str) -> int | None:
        match = re.search(rf"\b{label}\s+([\d,]+)", text, re.I)
        return parse_num(match.group(1)) if match else None

    return {
        "likes": find("L") or 0,
        "comments": None if platform == "YouTube" else (find("C") or 0),
        "saves": find("Saves") if platform == "TikTok" else None,
        "shares": find("Shares") if platform == "TikTok" else None,
    }


def load_supplement_docx(path: Path | None, start: date, end: date) -> list[dict[str, Any]]:
    if not path or not path.exists():
        return []
    doc = Document(path)
    posts = []
    current_creator = None
    current_platform = None
    rank_re = re.compile(r"^\d+\.\s+(.+)$")
    for block in iter_doc_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            match = rank_re.match(text)
            if match:
                current_creator = canonical_creator(match.group(1))
                current_platform = None
            elif text in {"Instagram", "TikTok", "YouTube"}:
                current_platform = text
            continue
        if not current_creator or current_platform not in {"Instagram", "TikTok", "YouTube"}:
            continue
        headers = [cell.text.strip() for cell in block.rows[0].cells]
        if headers[:6] != ["#", "Date", "Post Reference", "Views", "Engagement", "Category"]:
            continue
        for row in block.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 6 or not cells[0]:
                continue
            try:
                published = datetime.strptime(cells[1], "%b %d, %Y").date()
            except ValueError:
                continue
            if not (start <= published <= end):
                continue
            title, handle, pid = parse_ref(cells[2])
            metrics = parse_engagement(cells[4], current_platform)
            posts.append(
                {
                    "creator": current_creator,
                    "platform": current_platform,
                    "handle": handle,
                    "date": published,
                    "post_type": "Reel" if current_platform == "Instagram" else ("Short" if current_platform == "YouTube" else "Video"),
                    "post_id": pid,
                    "title": title,
                    "views": parse_num(cells[3]) or 0,
                    "category": cells[5] or category_for(title),
                    "url": post_url(current_platform, handle, pid),
                    "_source": "supplement_docx",
                    **metrics,
                }
            )
    return posts


def find_chrome(chrome_path: str | None = None) -> str | None:
    candidates = [
        chrome_path,
        os.environ.get("CHROME_BIN"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        shutil.which("google-chrome"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def fetch_dom(url: str, chrome_path: str | None = None, timeout: int = 35) -> str:
    chrome = find_chrome(chrome_path)
    if not chrome:
        return ""
    try:
        cp = subprocess.run(
            [chrome, "--headless=new", "--disable-gpu", "--no-sandbox", "--dump-dom", url],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            timeout=timeout,
        )
        return cp.stdout or ""
    except Exception:
        return ""


def scrape_instagram_profile(handle: str, chrome_path: str | None = None) -> tuple[str | None, dict[str, Any]]:
    profile = {"handle": handle, "verified": "N/A"}
    text = fetch_dom(f"https://www.instagram.com/{handle}/", chrome_path=chrome_path)
    if not text:
        try:
            text = requests.get(f"https://www.instagram.com/{handle}/", headers=HEADERS, timeout=20).text
        except Exception:
            text = ""
    ids = []
    for pattern in [r'"props":\{"id":"(\d+)"', r'"profile_id":"(\d+)"', r"profilePage_(\d+)", r'"container_id":"(\d+)"']:
        ids += re.findall(pattern, text)
    title = re.search(r"<title>(.*?)</title>", text, re.S)
    desc = re.search(r'<meta content="([^"]+)" name="description"', text)
    desc_text = html.unescape(desc.group(1)) if desc else ""
    if title:
        title_text = html.unescape(title.group(1))
        match = re.match(r"(.+?)\s*\(@" + re.escape(handle) + r"\)", title_text, re.I)
        if match:
            profile["full_name"] = clean_text(match.group(1), 80)
    count_match = re.search(r"([\d,.KMkm]+)\s+followers,\s+([\d,.KMkm]+)\s+following,\s+([\d,.KMkm]+)\s+posts", desc_text)
    if count_match:
        profile["followers"] = parse_num(count_match.group(1))
        profile["following"] = parse_num(count_match.group(2))
        profile["posts"] = parse_num(count_match.group(3))
    bio_match = re.search(r'on Instagram:\s+"(.*?)"', desc_text)
    if bio_match:
        profile["bio"] = clean_text(bio_match.group(1), 160)
    return (ids[0] if ids else None), profile


def scrape_instagram_posts(creator: str, handle: str, start: date, end: date, chrome_path: str | None = None) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    user_id, profile = scrape_instagram_profile(handle, chrome_path)
    posts = []
    if not user_id:
        return {"creator": creator, **profile}, posts
    session = requests.Session()
    session.headers.update(IG_HEADERS)
    max_id = None
    seen = set()
    for _ in range(8):
        params = {"count": "24"}
        if max_id:
            params["max_id"] = max_id
        try:
            response = session.get(f"https://www.instagram.com/api/v1/feed/user/{user_id}/", params=params, timeout=20)
            if response.status_code != 200:
                break
            data = response.json()
        except Exception:
            break
        for item in data.get("items", []):
            code = item.get("code") or item.get("shortcode")
            if not code or code in seen:
                continue
            seen.add(code)
            taken_at = item.get("taken_at") or item.get("device_timestamp")
            try:
                published = datetime.fromtimestamp(int(taken_at)).date()
            except Exception:
                continue
            if published < start:
                return {"creator": creator, **profile}, posts
            if not (start <= published <= end):
                continue
            caption = ((item.get("caption") or {}).get("text") or "")
            views = item.get("play_count") or item.get("video_view_count") or item.get("view_count") or 0
            posts.append(
                {
                    "creator": creator,
                    "platform": "Instagram",
                    "handle": handle,
                    "date": published,
                    "post_type": "Reel",
                    "post_id": code,
                    "title": clean_text(caption or code, 220),
                    "views": views,
                    "likes": item.get("like_count") or 0,
                    "comments": item.get("comment_count") or 0,
                    "saves": None,
                    "shares": None,
                    "category": category_for(caption),
                    "url": post_url("Instagram", handle, code),
                    "_source": "instagram_public_feed",
                }
            )
        max_id = data.get("next_max_id")
        if not max_id:
            break
        time.sleep(0.4)
    return {"creator": creator, **profile}, posts


def scrape_tiktok(creator: str, handle: str, start: date, end: date) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    profile = {"creator": creator, "handle": handle, "verified": "N/A"}
    posts = []
    try:
        response = requests.get("https://www.tikwm.com/api/user/info", params={"unique_id": handle}, headers=HEADERS, timeout=20)
        data = response.json().get("data") or {}
        user = data.get("user") or data.get("userInfo", {}).get("user") or {}
        stats = data.get("stats") or data.get("userInfo", {}).get("stats") or {}
        profile.update(
            {
                "followers": stats.get("followerCount"),
                "total_likes": stats.get("heartCount"),
                "videos": stats.get("videoCount"),
                "verified": "Yes" if user.get("verified") else "No",
                "bio": clean_text(user.get("signature"), 160),
            }
        )
    except Exception:
        pass
    cursor = "0"
    for _ in range(6):
        try:
            response = requests.get("https://www.tikwm.com/api/user/posts", params={"unique_id": handle, "count": 35, "cursor": cursor}, headers=HEADERS, timeout=20)
            data = response.json().get("data") or {}
            videos = data.get("videos") or data.get("list") or []
        except Exception:
            break
        if not videos:
            break
        for item in videos:
            pid = str(item.get("video_id") or item.get("id") or "")
            create_time = item.get("create_time") or item.get("createTime")
            try:
                published = datetime.fromtimestamp(int(create_time)).date()
            except Exception:
                continue
            if published < start:
                return profile, posts
            if not (start <= published <= end):
                continue
            title = clean_text(item.get("title") or item.get("desc") or pid, 220)
            posts.append(
                {
                    "creator": creator,
                    "platform": "TikTok",
                    "handle": handle,
                    "date": published,
                    "post_type": "Video",
                    "post_id": pid,
                    "title": title,
                    "views": item.get("play_count") or item.get("playCount") or 0,
                    "likes": item.get("digg_count") or item.get("diggCount") or 0,
                    "comments": item.get("comment_count") or item.get("commentCount") or 0,
                    "saves": item.get("collect_count") or item.get("collectCount") or 0,
                    "shares": item.get("share_count") or item.get("shareCount") or 0,
                    "category": category_for(title),
                    "url": post_url("TikTok", handle, pid),
                    "_source": "tikwm_public_api",
                }
            )
        cursor = str(data.get("cursor") or data.get("maxCursor") or "")
        if not cursor:
            break
        time.sleep(0.3)
    return profile, posts


def resolve_youtube_channel(handle: str) -> dict[str, Any]:
    profile = {"handle": handle}
    try:
        response = requests.get(f"https://www.youtube.com/@{handle}", headers=HEADERS, timeout=20)
        text = response.text
    except Exception:
        return profile
    patterns = [r'"channelId":"(UC[A-Za-z0-9_-]+)"', r'youtube\.com/channel/(UC[A-Za-z0-9_-]+)']
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            profile["channel_id"] = match.group(1)
            break
    subs = re.search(r'"subscriberCountText":\{"simpleText":"([\d,.KMkm]+)\s+subscribers?"', text)
    views = re.search(r'"viewCountText":\{"simpleText":"([\d,.KMkm]+)\s+views?"', text)
    videos = re.search(r'"videosCountText":\{"runs":\[\{"text":"([\d,.KMkm]+)"', text)
    if subs:
        profile["subscribers"] = parse_num(subs.group(1))
    if views:
        profile["channel_views"] = parse_num(views.group(1))
    if videos:
        profile["profile_videos"] = parse_num(videos.group(1))
    return profile


def fetch_youtube_video_stats(video_id: str) -> dict[str, Any]:
    try:
        text = requests.get(f"https://www.youtube.com/shorts/{video_id}", headers=HEADERS, timeout=20).text
    except Exception:
        return {}
    stats = {}
    view_match = re.search(r'"viewCount":"(\d+)"', text) or re.search(r'"views":\{"simpleText":"([\d,.KMkm]+)\s+views?"', text)
    like_match = re.search(r'"likeCount":\s*"?(\d+)"?', text)
    if view_match:
        stats["views"] = parse_num(view_match.group(1)) or 0
    if like_match:
        stats["likes"] = parse_num(like_match.group(1)) or 0
    return stats


def scrape_youtube(creator: str, handle: str, start: date, end: date) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    profile = {"creator": creator, **resolve_youtube_channel(handle)}
    posts = []
    channel_id = profile.get("channel_id")
    if not channel_id:
        return profile, posts
    try:
        rss = requests.get("https://www.youtube.com/feeds/videos.xml", params={"channel_id": channel_id}, headers=HEADERS, timeout=20)
        root = ET.fromstring(rss.text)
    except Exception:
        return profile, posts
    ns = {"atom": "http://www.w3.org/2005/Atom", "yt": "http://www.youtube.com/xml/schemas/2015"}
    for entry in root.findall("atom:entry", ns):
        vid = entry.findtext("yt:videoId", default="", namespaces=ns)
        title = clean_text(entry.findtext("atom:title", default=vid, namespaces=ns), 220)
        published_raw = entry.findtext("atom:published", default="", namespaces=ns)
        try:
            published = datetime.fromisoformat(published_raw.replace("Z", "+00:00")).date()
        except Exception:
            continue
        if not (start <= published <= end):
            continue
        stats = fetch_youtube_video_stats(vid)
        posts.append(
            {
                "creator": creator,
                "platform": "YouTube",
                "handle": handle,
                "date": published,
                "post_type": "Short",
                "post_id": vid,
                "title": title,
                "views": stats.get("views", 0),
                "likes": stats.get("likes", 0),
                "comments": None,
                "saves": None,
                "shares": None,
                "category": category_for(title),
                "url": post_url("YouTube", handle, vid),
                "_source": "youtube_rss_public_page",
            }
        )
    return profile, posts


def scrape_public_accounts(roster: dict[str, dict[str, list[str]]], start: date, end: date, chrome_path: str | None = None) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    profiles_tt, profiles_ig, profiles_yt, posts = [], [], [], []
    tasks = []
    for creator, handles in roster.items():
        for handle in handles.get("ig", []):
            tasks.append(("ig", creator, handle))
        for handle in handles.get("tt", []):
            tasks.append(("tt", creator, handle))
        for handle in handles.get("yt", []):
            tasks.append(("yt", creator, handle))
    if not tasks:
        return profiles_tt, profiles_ig, profiles_yt, posts

    def run_task(task):
        platform, creator, handle = task
        if platform == "ig":
            profile, found = scrape_instagram_posts(creator, handle, start, end, chrome_path)
        elif platform == "tt":
            profile, found = scrape_tiktok(creator, handle, start, end)
        else:
            profile, found = scrape_youtube(creator, handle, start, end)
        return platform, profile, found

    with ThreadPoolExecutor(max_workers=min(8, len(tasks))) as executor:
        futures = [executor.submit(run_task, task) for task in tasks]
        for future in as_completed(futures):
            try:
                platform, profile, found = future.result()
            except Exception:
                continue
            if platform == "ig":
                profiles_ig.append(profile)
            elif platform == "tt":
                profiles_tt.append(profile)
            else:
                profiles_yt.append(profile)
            posts.extend(found)
    return profiles_tt, profiles_ig, profiles_yt, posts


def extract_meta_description(page_html: str) -> str:
    for tag in re.findall(r"<meta\b[^>]*>", page_html, flags=re.I | re.S):
        if not (re.search(r'\bname=["\']description["\']', tag, re.I) or re.search(r'\bproperty=["\']og:description["\']', tag, re.I)):
            continue
        content = re.search(r'\bcontent=["\']([^"\']*)["\']', tag, re.I | re.S)
        if content:
            return content.group(1)
    return ""


def parse_instagram_meta(description: str) -> dict[str, Any]:
    desc = html.unescape(description or "").replace("\xa0", " ")
    match = re.search(
        r"([\d,]+)\s+likes?,\s+([\d,]+)\s+comments?\s+-\s+(.+?)\s+on\s+([A-Za-z]+ \d{1,2}, \d{4}):\s+\"(.*?)(?:\"\.\s*)?$",
        desc,
        re.S,
    )
    if not match:
        return {}
    return {
        "likes": parse_num(match.group(1)) or 0,
        "comments": parse_num(match.group(2)) or 0,
        "handle": match.group(3).strip().lstrip("@"),
        "date": datetime.strptime(match.group(4), "%B %d, %Y").date(),
        "title": clean_text(match.group(5).rstrip(". "), 220),
    }


def fetch_instagram_page_meta(url: str, chrome_path: str | None = None) -> dict[str, Any]:
    try:
        response = requests.get(url, headers=HEADERS, timeout=15)
        record = parse_instagram_meta(extract_meta_description(response.text))
        if record:
            return record
    except Exception:
        pass
    text = fetch_dom(url, chrome_path=chrome_path)
    return parse_instagram_meta(extract_meta_description(text))


def verify_instagram_posts(posts: list[dict[str, Any]], chrome_path: str | None = None, cache_path: Path | None = None) -> dict[str, int]:
    if cache_path and cache_path.exists():
        cache = json.loads(cache_path.read_text(encoding="utf-8"))
    else:
        cache = {}
    session = requests.Session()
    session.headers.update(HEADERS)
    checked = 0
    refreshed = 0
    missing_meta = []
    for post in posts:
        if post["platform"] != "Instagram":
            continue
        checked += 1
        key = post.get("post_id") or post.get("url")
        record = cache.get(key) or {}
        if not all(record.get(field) for field in ["title", "handle", "media_id"]):
            try:
                response = session.get("https://www.instagram.com/api/v1/oembed/", params={"url": post["url"]}, timeout=15)
                if response.ok:
                    data = response.json()
                    record["title"] = clean_text(data.get("title") or "", 220)
                    record["handle"] = (data.get("author_name") or "").strip().lstrip("@")
                    record["media_id"] = data.get("media_id")
            except Exception:
                pass
            cache[key] = record
            time.sleep(0.05)
        if not all(field in record for field in ["likes", "comments", "date"]):
            missing_meta.append((key, post["url"]))
    if missing_meta:
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = {executor.submit(fetch_instagram_page_meta, url, chrome_path): key for key, url in missing_meta}
            for future in as_completed(futures):
                key = futures[future]
                record = future.result() or {}
                if record:
                    cache.setdefault(key, {}).update(record)
    for post in posts:
        if post["platform"] != "Instagram":
            continue
        key = post.get("post_id") or post.get("url")
        record = cache.get(key) or {}
        for field in ["title", "handle", "likes", "comments", "date"]:
            if field in record and record[field] not in [None, ""]:
                value = record[field]
                if field == "date" and isinstance(value, str):
                    value = parse_date(value)
                if post.get(field) != value:
                    refreshed += 1
                    post[field] = value
    if cache_path:
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        cache_path.write_text(json.dumps(cache, default=str, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"checked": checked, "refreshed_fields": refreshed}


def dedupe_posts(posts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: dict[tuple[str, str], dict[str, Any]] = {}
    source_rank = {"csv": 0, "supplement_docx": 1, "instagram_public_feed": 2, "tikwm_public_api": 2, "youtube_rss_public_page": 2}
    for post in posts:
        if not post.get("creator") or not post.get("platform"):
            continue
        key = (post["platform"], post.get("post_id") or post.get("url") or f"{post['creator']}:{post['date']}:{post['title']}")
        existing = merged.get(key)
        if not existing:
            merged[key] = post
            continue
        old_rank = source_rank.get(existing.get("_source"), 0)
        new_rank = source_rank.get(post.get("_source"), 0)
        if new_rank >= old_rank:
            combined = {**existing, **{k: v for k, v in post.items() if v not in [None, "", "N/A"]}}
            merged[key] = combined
    return list(merged.values())


def summarize(posts: list[dict[str, Any]]) -> dict[str, int]:
    return {
        "posts": len(posts),
        "views": sum(int(p.get("views") or 0) for p in posts),
        "likes": sum(int(p.get("likes") or 0) for p in posts),
        "comments": sum(int(p.get("comments") or 0) for p in posts if p.get("comments") is not None),
        "saves": sum(int(p.get("saves") or 0) for p in posts if p.get("saves") is not None),
        "shares": sum(int(p.get("shares") or 0) for p in posts if p.get("shares") is not None),
    }


def build_creator_summaries(roster, posts, tt_profiles, ig_profiles, yt_profiles):
    by_creator = defaultdict(list)
    for post in posts:
        by_creator[post["creator"]].append(post)
    tt_by, ig_by, yt_by = defaultdict(list), defaultdict(list), defaultdict(list)
    for row in tt_profiles:
        tt_by[row["creator"]].append(row)
    for row in ig_profiles:
        ig_by[row["creator"]].append(row)
    for row in yt_profiles:
        yt_by[row["creator"]].append(row)
    rows = []
    for creator, creator_posts in by_creator.items():
        platform_posts = {p: [x for x in creator_posts if x["platform"] == p] for p in ["Instagram", "TikTok", "YouTube"]}
        ig, tt, yt = summarize(platform_posts["Instagram"]), summarize(platform_posts["TikTok"]), summarize(platform_posts["YouTube"])
        tt_followers = sum(x.get("followers") or 0 for x in tt_by[creator]) or None
        tt_total_likes = sum(x.get("total_likes") or 0 for x in tt_by[creator]) or None
        tt_videos = sum(x.get("videos") or 0 for x in tt_by[creator]) or None
        ig_followers = sum(x.get("followers") or 0 for x in ig_by[creator]) or None
        ig_posts = sum(x.get("posts") or 0 for x in ig_by[creator]) or None
        yt_subs = sum(x.get("subscribers") or 0 for x in yt_by[creator]) or None
        yt_profile_videos = sum(x.get("profile_videos") or 0 for x in yt_by[creator]) or None
        top_platform = max([("Instagram", ig["views"]), ("TikTok", tt["views"]), ("YouTube", yt["views"])], key=lambda x: x[1])[0]
        rows.append(
            {
                "creator": creator,
                "handles": " | ".join(
                    [
                        "TT: " + ", ".join("@" + h for h in roster.get(creator, {}).get("tt", [])),
                        "IG: " + ", ".join("@" + h for h in roster.get(creator, {}).get("ig", [])),
                        "YT: " + ", ".join("@" + h for h in roster.get(creator, {}).get("yt", [])),
                    ]
                ),
                "tt_followers": tt_followers,
                "tt_total_likes": tt_total_likes,
                "tt_videos": tt_videos,
                "ig_followers": ig_followers,
                "ig_posts": ig_posts,
                "ig_agg_likes": ig["likes"],
                "ig_agg_comments": ig["comments"],
                "ig_agg_views": ig["views"],
                "ig_avg_likes": safe_div(ig["likes"], ig["posts"]),
                "ig_eng_rate": safe_div(ig["likes"] + ig["comments"], ig["views"]) * 100 if ig["views"] else None,
                "yt_subscribers": yt_subs,
                "yt_total_views": yt["views"],
                "yt_total_likes": yt["likes"],
                "yt_videos": yt_profile_videos or yt["posts"],
                "yt_avg_views": safe_div(yt["views"], yt["posts"]),
                "yt_avg_likes": safe_div(yt["likes"], yt["posts"]),
                "yt_eng_rate": safe_div(yt["likes"], yt["views"]) * 100 if yt["views"] else None,
                "total_audience": sum(x for x in [tt_followers or 0, ig_followers or 0, yt_subs or 0]),
                "top_platform": top_platform,
            }
        )
    return sorted(rows, key=lambda r: r["ig_agg_views"] + r["yt_total_views"], reverse=True)


def add_hyperlink(paragraph, text, url, color="0563C1", underline=True):
    if not url:
        paragraph.add_run(str(text))
        return
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), color)
    props.append(color_elem)
    if underline:
        underline_elem = OxmlElement("w:u")
        underline_elem.set(qn("w:val"), "single")
        props.append(underline_elem)
    run.append(props)
    text_elem = OxmlElement("w:t")
    text_elem.text = str(text)
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(18)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(10)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text="", bold=False, size=7.0):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    return paragraph


def table_from_rows(doc, headers, rows, fill="D9EAF7", size=7.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True, size=size)
        shade_cell(cell, fill)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = set_cell_text(cells[idx], "", size=size)
            if isinstance(value, tuple):
                add_hyperlink(paragraph, value[0], value[1])
            else:
                paragraph.add_run(str(value)).font.size = Pt(size)
    return table


def build_docx(path: Path, campaign_name: str, start: date, end: date, roster, posts, summaries, tt_profiles, ig_profiles, yt_profiles):
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{campaign_name} - Creator Analytics")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{fmt_date(start)} - {fmt_date(end)} | Instagram, TikTok and YouTube").bold = True
    count = doc.add_paragraph()
    count.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count.add_run(f"{fmt_num(len(posts))} posts | {len(set(p['creator'] for p in posts))} active creators").bold = True

    doc.add_heading("Cross-Platform Summary", level=1)
    summary_headers = [
        "Creator Name", "Handles", "TT Followers", "TT Total Likes", "TT Videos", "IG Followers", "IG Posts",
        "IG Agg Likes", "IG Agg Comments", "IG Agg Views", "IG Avg Likes", "IG Eng Rate %",
        "YT Subscribers", "YT Total Views", "YT Total Likes", "YT Total Comments", "YT Videos", "YT Avg Views",
        "YT Avg Likes", "YT Eng Rate %", "Total Audience", "Top Platform",
    ]
    summary_rows = []
    for r in summaries:
        summary_rows.append(
            [
                DISPLAY_NAMES.get(r["creator"], r["creator"]), r["handles"], fmt_num(r["tt_followers"]), fmt_num(r["tt_total_likes"]),
                fmt_num(r["tt_videos"]), fmt_num(r["ig_followers"]), fmt_num(r["ig_posts"]), fmt_num(r["ig_agg_likes"]),
                fmt_num(r["ig_agg_comments"]), fmt_num(r["ig_agg_views"]), "N/A" if r["ig_avg_likes"] is None else f"{r['ig_avg_likes']:.1f}",
                fmt_pct(r["ig_eng_rate"]), fmt_num(r["yt_subscribers"]), fmt_num(r["yt_total_views"]), fmt_num(r["yt_total_likes"]),
                "N/A", fmt_num(r["yt_videos"]), "N/A" if r["yt_avg_views"] is None else f"{r['yt_avg_views']:.0f}",
                "N/A" if r["yt_avg_likes"] is None else f"{r['yt_avg_likes']:.1f}", fmt_pct(r["yt_eng_rate"]),
                fmt_num(r["total_audience"]), r["top_platform"],
            ]
        )
    table_from_rows(doc, summary_headers, summary_rows, size=5.8)

    doc.add_heading("Platform Analytics Snapshot", level=1)
    platform_rows = []
    for platform in ["Instagram", "TikTok", "YouTube"]:
        rows = [p for p in posts if p["platform"] == platform]
        s = summarize(rows)
        vals = [p.get("views") or 0 for p in rows]
        platform_rows.append(
            [
                platform, fmt_num(s["posts"]), fmt_num(s["views"]), fmt_num(round(safe_div(s["views"], s["posts"]) or 0)),
                fmt_num(median(vals) if vals else 0), fmt_num(s["likes"]), "N/A" if platform == "YouTube" else fmt_num(s["comments"]),
                "N/A / N/A" if platform != "TikTok" else f"{fmt_num(s['saves'])} / {fmt_num(s['shares'])}",
            ]
        )
    table_from_rows(doc, ["Platform", "Posts", "Views/Plays", "Avg", "Median", "Likes", "Comments", "Saves/Shares"], platform_rows, fill="EAF4E1")

    doc.add_heading("Creator Ranking", level=1)
    rank_rows = []
    by_creator = defaultdict(list)
    for post in posts:
        by_creator[post["creator"]].append(post)
    ranked = sorted(by_creator.items(), key=lambda item: sum(p.get("views") or 0 for p in item[1]), reverse=True)
    for idx, (creator, cposts) in enumerate(ranked, start=1):
        top = max(cposts, key=lambda p: p.get("views") or 0)
        rank_rows.append(
            [
                idx, DISPLAY_NAMES.get(creator, creator), len(cposts),
                sum(1 for p in cposts if p["platform"] == "Instagram"),
                sum(1 for p in cposts if p["platform"] == "TikTok"),
                sum(1 for p in cposts if p["platform"] == "YouTube"),
                fmt_num(sum(p.get("views") or 0 for p in cposts)),
                f"{top['platform']} | {fmt_num(top.get('views'))}",
            ]
        )
    table_from_rows(doc, ["Rank", "Creator", "Posts", "IG", "TikTok", "YT", "Views/Plays", "Top Post"], rank_rows, fill="EAF4E1")

    doc.add_heading("Creator Post References", level=1)
    for idx, (creator, cposts) in enumerate(ranked, start=1):
        if idx > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_heading(f"{idx}. {DISPLAY_NAMES.get(creator, creator)}", level=1)
        doc.add_paragraph(
            f"Summary: {len(cposts)} posts | {fmt_num(sum(p.get('views') or 0 for p in cposts))} views/plays | "
            f"IG {sum(1 for p in cposts if p['platform'] == 'Instagram')}, "
            f"TikTok {sum(1 for p in cposts if p['platform'] == 'TikTok')}, "
            f"YouTube {sum(1 for p in cposts if p['platform'] == 'YouTube')}"
        )
        for platform in ["Instagram", "TikTok", "YouTube"]:
            rows = sorted([p for p in cposts if p["platform"] == platform], key=lambda p: (p.get("views") or 0, p["date"]), reverse=True)
            doc.add_heading(platform, level=2)
            if not rows:
                doc.add_paragraph("0 posts.")
                continue
            post_rows = []
            for n, p in enumerate(rows, start=1):
                post_rows.append(
                    [
                        n, fmt_date(p["date"]), (clean_text(p.get("title") or "Open post", 140), p.get("url")),
                        fmt_num(p.get("views")), fmt_num(p.get("likes")),
                        "N/A" if p.get("comments") is None else fmt_num(p.get("comments")),
                        "N/A" if p.get("saves") is None else fmt_num(p.get("saves")),
                        "N/A" if p.get("shares") is None else fmt_num(p.get("shares")),
                        p.get("category") or category_for(p.get("title")),
                    ]
                )
            table_from_rows(doc, ["#", "Date", "Post Reference", "Views", "Likes", "Comments", "Saves", "Shares", "Category"], post_rows, fill="F2F2F2", size=6.8)
    doc.save(path)


def write_sheet(workbook, name, headers, rows):
    ws = workbook.add_worksheet(name[:31])
    header_fmt = workbook.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1, "text_wrap": True})
    text_fmt = workbook.add_format({"border": 1, "text_wrap": True, "valign": "top"})
    num_fmt = workbook.add_format({"border": 1, "num_format": "#,##0"})
    url_fmt = workbook.add_format({"font_color": "blue", "underline": 1, "border": 1, "text_wrap": True})
    for c, header in enumerate(headers):
        ws.write(0, c, header, header_fmt)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            if isinstance(value, tuple):
                text, url = value
                if url:
                    ws.write_url(r, c, url, url_fmt, string=text or url)
                else:
                    ws.write(r, c, text, text_fmt)
            elif isinstance(value, (int, float)) and not isinstance(value, bool):
                ws.write_number(r, c, value, num_fmt)
            else:
                ws.write(r, c, value, text_fmt)
    ws.freeze_panes(1, 0)
    ws.autofilter(0, 0, max(len(rows), 1), len(headers) - 1)
    for c, header in enumerate(headers):
        ws.set_column(c, c, min(max(len(header) + 2, 12), 42))


def build_xlsx(path: Path, summaries, posts):
    workbook = xlsxwriter.Workbook(path)
    summary_headers = [
        "Creator Name", "Handles", "TT Followers", "TT Total Likes", "TT Videos", "IG Followers", "IG Posts",
        "IG Agg Likes", "IG Agg Comments", "IG Agg Views", "IG Avg Likes", "IG Eng Rate %",
        "YT Subscribers", "YT Total Views", "YT Total Likes", "YT Total Comments", "YT Videos", "YT Avg Views",
        "YT Avg Likes", "YT Eng Rate %", "Total Audience", "Top Platform",
    ]
    write_sheet(workbook, "Cross Platform Summary", summary_headers, [
        [
            DISPLAY_NAMES.get(r["creator"], r["creator"]), r["handles"], r["tt_followers"], r["tt_total_likes"], r["tt_videos"],
            r["ig_followers"], r["ig_posts"], r["ig_agg_likes"], r["ig_agg_comments"], r["ig_agg_views"],
            r["ig_avg_likes"], r["ig_eng_rate"], r["yt_subscribers"], r["yt_total_views"], r["yt_total_likes"],
            "N/A", r["yt_videos"], r["yt_avg_views"], r["yt_avg_likes"], r["yt_eng_rate"], r["total_audience"], r["top_platform"],
        ]
        for r in summaries
    ])
    platform_rows = []
    for platform in ["Instagram", "TikTok", "YouTube"]:
        rows = [p for p in posts if p["platform"] == platform]
        s = summarize(rows)
        vals = [p.get("views") or 0 for p in rows]
        platform_rows.append([platform, s["posts"], s["views"], round(safe_div(s["views"], s["posts"]) or 0), median(vals) if vals else 0, s["likes"], "N/A" if platform == "YouTube" else s["comments"], "N/A / N/A" if platform != "TikTok" else f"{s['saves']} / {s['shares']}"])
    write_sheet(workbook, "Platform Snapshot", ["Platform", "Posts", "Views/Plays", "Avg", "Median", "Likes", "Comments", "Saves/Shares"], platform_rows)
    for platform in ["Instagram", "TikTok", "YouTube"]:
        rows = []
        for p in sorted([x for x in posts if x["platform"] == platform], key=lambda x: (x["creator"].lower(), -(x.get("views") or 0))):
            rows.append(
                [
                    DISPLAY_NAMES.get(p["creator"], p["creator"]), p["handle"], p["date"].isoformat(), p["post_type"], p["title"],
                    p["views"], p["likes"], "N/A" if p.get("comments") is None else p.get("comments"),
                    "N/A" if p.get("saves") is None else p.get("saves"),
                    "N/A" if p.get("shares") is None else p.get("shares"),
                    p.get("category") or category_for(p.get("title")), (p["url"], p["url"]),
                ]
            )
        write_sheet(workbook, f"{platform} Posts", ["Creator", "Handle", "Date", "Post Type", "Caption/Title", "Views", "Likes", "Comments", "Saves", "Shares", "Category", "URL"], rows)
    workbook.close()


def build_csv_exports(output_dir: Path, stem: str, summaries, posts) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = []

    posts_path = output_dir / f"{stem}_posts.csv"
    with posts_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "creator", "platform", "handle", "date", "post_type", "post_id", "title", "views",
                "likes", "comments", "saves", "shares", "category", "url", "source",
            ],
        )
        writer.writeheader()
        for post in sorted(posts, key=lambda p: (p["creator"].lower(), p["platform"], p["date"], -(p.get("views") or 0))):
            writer.writerow(
                {
                    "creator": DISPLAY_NAMES.get(post["creator"], post["creator"]),
                    "platform": post.get("platform"),
                    "handle": post.get("handle"),
                    "date": post.get("date").isoformat() if post.get("date") else "",
                    "post_type": post.get("post_type"),
                    "post_id": post.get("post_id"),
                    "title": post.get("title"),
                    "views": post.get("views"),
                    "likes": post.get("likes"),
                    "comments": "" if post.get("comments") is None else post.get("comments"),
                    "saves": "" if post.get("saves") is None else post.get("saves"),
                    "shares": "" if post.get("shares") is None else post.get("shares"),
                    "category": post.get("category"),
                    "url": post.get("url"),
                    "source": post.get("_source", "public_scrape"),
                }
            )
    paths.append(posts_path)

    summary_path = output_dir / f"{stem}_creator_summary.csv"
    with summary_path.open("w", newline="", encoding="utf-8-sig") as f:
        headers = [
            "creator", "handles", "tt_followers", "tt_total_likes", "tt_videos", "ig_followers", "ig_posts",
            "ig_agg_likes", "ig_agg_comments", "ig_agg_views", "ig_avg_likes", "ig_eng_rate",
            "yt_subscribers", "yt_total_views", "yt_total_likes", "yt_videos", "yt_avg_views",
            "yt_avg_likes", "yt_eng_rate", "total_audience", "top_platform",
        ]
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        for row in summaries:
            writer.writerow({**row, "creator": DISPLAY_NAMES.get(row["creator"], row["creator"])})
    paths.append(summary_path)

    platform_path = output_dir / f"{stem}_platform_summary.csv"
    with platform_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=["platform", "posts", "views", "avg_views", "median_views", "likes", "comments", "saves", "shares"])
        writer.writeheader()
        for platform in ["Instagram", "TikTok", "YouTube"]:
            rows = [p for p in posts if p["platform"] == platform]
            s = summarize(rows)
            vals = [p.get("views") or 0 for p in rows]
            writer.writerow(
                {
                    "platform": platform,
                    "posts": s["posts"],
                    "views": s["views"],
                    "avg_views": round(safe_div(s["views"], s["posts"]) or 0),
                    "median_views": median(vals) if vals else 0,
                    "likes": s["likes"],
                    "comments": "" if platform == "YouTube" else s["comments"],
                    "saves": "" if platform != "TikTok" else s["saves"],
                    "shares": "" if platform != "TikTok" else s["shares"],
                }
            )
    paths.append(platform_path)
    return paths


def collect_data(config: RunConfig) -> dict[str, Any]:
    config.output_dir.mkdir(parents=True, exist_ok=True)
    roster = load_creator_registry(config.roster_json, config.creators_csv, config.google_sheet_csv_url)
    profiles_tt, profiles_ig, profiles_yt = [], [], []
    all_posts = []
    if not roster:
        raise ValueError("No creator registry found. Add roster_json, creators_csv, or google_sheet_csv_url.")
    if config.scrape_public:
        tt, ig, yt, public_posts = scrape_public_accounts(roster, config.start_date, config.end_date, config.chrome_path)
        profiles_tt.extend(tt)
        profiles_ig.extend(ig)
        profiles_yt.extend(yt)
        all_posts.extend(public_posts)
    for supplemental_csv in config.supplemental_analytics_csvs or []:
        all_posts.extend(load_csv_posts(supplemental_csv, config.start_date, config.end_date))
    all_posts.extend(load_supplement_docx(config.supplemental_docx, config.start_date, config.end_date))
    posts = dedupe_posts(all_posts)
    if config.verify_instagram:
        verify_instagram_posts(posts, config.chrome_path, config.output_dir / ".ig_verify_cache.json")
    summaries = build_creator_summaries(roster, posts, profiles_tt, profiles_ig, profiles_yt)
    return {
        "roster": roster,
        "profiles_tt": profiles_tt,
        "profiles_ig": profiles_ig,
        "profiles_yt": profiles_yt,
        "posts": posts,
        "summaries": summaries,
    }


def serialize_payload(data: dict[str, Any], config: RunConfig) -> dict[str, Any]:
    def convert(value):
        if isinstance(value, date):
            return value.isoformat()
        if isinstance(value, Path):
            return str(value)
        if isinstance(value, list):
            return [convert(v) for v in value]
        if isinstance(value, dict):
            return {k: convert(v) for k, v in value.items()}
        return value

    return {
        "campaign_name": config.campaign_name,
        "start_date": config.start_date.isoformat(),
        "end_date": config.end_date.isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z"),
        "platform_counts": {platform: sum(1 for p in data["posts"] if p["platform"] == platform) for platform in ["Instagram", "TikTok", "YouTube"]},
        "posts": convert(data["posts"]),
        "summaries": convert(data["summaries"]),
        "roster": convert(data["roster"]),
    }


def deserialize_payload(payload: dict[str, Any]) -> dict[str, Any]:
    posts = []
    for post in payload.get("posts", []):
        row = dict(post)
        if row.get("date"):
            row["date"] = parse_date(row["date"])
        posts.append(row)
    return {
        "posts": posts,
        "summaries": payload.get("summaries", []),
        "roster": payload.get("roster", {}),
    }


def write_snapshot(path: Path, data: dict[str, Any], config: RunConfig) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = serialize_payload(data, config)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def load_snapshot(path: Path) -> dict[str, Any] | None:
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def build_exports_from_payload(
    payload: dict[str, Any],
    output_dir: Path,
    output_stem: str | None = None,
    export_csv: bool = True,
    export_xlsx: bool = True,
    export_docx: bool = True,
) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    restored = deserialize_payload(payload)
    posts = restored["posts"]
    summaries = restored["summaries"]
    roster = restored["roster"]
    campaign_name = payload.get("campaign_name", "Zensi")
    start = parse_date(payload.get("start_date"))
    end = parse_date(payload.get("end_date"))
    stem = output_stem or f"{campaign_name.replace(' ', '_')}_CreatorAnalytics_{start}_to_{end}"
    docx_path = output_dir / f"{stem}.docx"
    xlsx_path = output_dir / f"{stem}.xlsx"
    csv_paths = []
    if export_docx:
        build_docx(docx_path, campaign_name, start, end, roster, posts, summaries, [], [], [])
    else:
        docx_path = None
    if export_xlsx:
        build_xlsx(xlsx_path, summaries, posts)
    else:
        xlsx_path = None
    if export_csv:
        csv_paths = build_csv_exports(output_dir, stem, summaries, posts)
    return {
        "docx": docx_path,
        "xlsx": xlsx_path,
        "csvs": csv_paths,
        "posts": len(posts),
        "platform_counts": {platform: sum(1 for p in posts if p["platform"] == platform) for platform in ["Instagram", "TikTok", "YouTube"]},
    }


def run_report(config: RunConfig) -> dict[str, Any]:
    config.output_dir.mkdir(parents=True, exist_ok=True)
    data = collect_data(config)
    posts = data["posts"]
    summaries = data["summaries"]
    roster = data["roster"]
    profiles_tt = data["profiles_tt"]
    profiles_ig = data["profiles_ig"]
    profiles_yt = data["profiles_yt"]
    stem = config.output_stem or f"{config.campaign_name.replace(' ', '_')}_CreatorAnalytics_{config.start_date}_to_{config.end_date}"
    docx_path = config.output_dir / f"{stem}.docx"
    xlsx_path = config.output_dir / f"{stem}.xlsx"
    csv_paths = []
    if config.export_docx:
        build_docx(docx_path, config.campaign_name, config.start_date, config.end_date, roster, posts, summaries, profiles_tt, profiles_ig, profiles_yt)
    else:
        docx_path = None
    if config.export_xlsx:
        build_xlsx(xlsx_path, summaries, posts)
    else:
        xlsx_path = None
    if config.export_csv:
        csv_paths = build_csv_exports(config.output_dir, stem, summaries, posts)
    return {
        "docx": docx_path,
        "xlsx": xlsx_path,
        "csvs": csv_paths,
        "posts": len(posts),
        "platform_counts": {platform: sum(1 for p in posts if p["platform"] == platform) for platform in ["Instagram", "TikTok", "YouTube"]},
    }
